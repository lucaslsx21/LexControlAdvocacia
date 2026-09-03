from io import BytesIO
from xml.sax.saxutils import escape
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import IntegrityError
from django.db.models.deletion import ProtectedError
from django.db.models import Q
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from .forms import AdvogadoForm, AndamentoForm, HistoricoForm, ImportarProcessoForm, MaterialForm, MovimentacaoForm, PeticaoIAForm, ProcessoForm, RevisarPeticaoForm
from .ia_peticoes import gerar_minuta_peticao
from .models import Advogado, Andamento, DocumentoProcesso, HistoricoAdvogado, Material, MovimentacaoEstoque, PeticaoIA, Processo
from .services import analisar_processo, extrair_texto

EXPORTACOES = {
    "advogados": ("Advogados", ["Nome", "OAB", "E-mail", "Admissão", "Ativo"], lambda x: [x.nome, x.oab, x.email, x.data_admissao.strftime("%d/%m/%Y"), "Sim" if x.ativo else "Não"], Advogado.objects.all),
    "processos": ("Processos", ["Número", "Título", "Área", "Status", "Advogado", "Atualização"], lambda x: [x.numero, x.titulo, x.get_area_display(), x.get_status_display(), x.advogado.nome, timezone.localtime(x.atualizado_em).strftime("%d/%m/%Y %H:%M")], lambda: Processo.objects.select_related("advogado").all()),
    "estoque": ("Estoque", ["Material", "Categoria", "Quantidade", "Unidade", "Estoque mínimo"], lambda x: [x.nome, x.categoria, x.quantidade, x.get_unidade_display(), x.estoque_minimo], Material.objects.all),
}

@login_required
def dashboard(request):
    return render(request, "juridico/dashboard.html", {
        "advogados": Advogado.objects.filter(ativo=True).count(), "processos": Processo.objects.count(),
        "deferidos": Processo.objects.filter(status="DEFERIDO").count(), "concluidos": Processo.objects.filter(status="CONCLUIDO").count(),
        "prazos": Processo.objects.filter(prazo_atual__gte=timezone.now()).order_by("prazo_atual")[:5],
        "estoque_baixo": Material.objects.filter(quantidade__lte=models.F("estoque_minimo"))[:5],
        "recentes": Processo.objects.select_related("advogado")[:6],
    })

@login_required
def advogados(request):
    q = request.GET.get("q", "")
    itens = Advogado.objects.filter(Q(nome__icontains=q) | Q(oab__icontains=q) | Q(especialidades__icontains=q))
    return render(request, "juridico/advogados.html", {"itens": itens, "q": q})

@login_required
def advogado_detalhe(request, pk):
    return render(request, "juridico/advogado_detalhe.html", {"item": get_object_or_404(Advogado, pk=pk)})

@login_required
def processos(request):
    itens = Processo.objects.select_related("advogado")
    q, area, status = request.GET.get("q", ""), request.GET.get("area", ""), request.GET.get("status", "")
    if q: itens = itens.filter(Q(numero__icontains=q) | Q(titulo__icontains=q) | Q(parte_autora__icontains=q) | Q(parte_re__icontains=q) | Q(descricao__icontains=q) | Q(advogado__nome__icontains=q))
    if area: itens = itens.filter(area=area)
    if status: itens = itens.filter(status=status)
    return render(request, "juridico/processos.html", {"itens": itens, "q": q, "area": area, "status": status, "areas": Processo.AREAS, "status_opcoes": Processo.STATUS})

@login_required
def processo_detalhe(request, pk):
    return render(request, "juridico/processo_detalhe.html", {"item": get_object_or_404(Processo.objects.select_related("advogado"), pk=pk)})

@login_required
def resultados(request):
    itens = Processo.objects.select_related("advogado").filter(status__in=["DEFERIDO", "CONCLUIDO"])
    return render(request, "juridico/resultados.html", {"itens": itens, "areas": Processo.AREAS})

@login_required
def busca_global(request):
    q = request.GET.get("q", "").strip()
    return render(request, "juridico/busca.html", {"q": q,
        "advogados": Advogado.objects.filter(Q(nome__icontains=q) | Q(oab__icontains=q) | Q(resumo_profissional__icontains=q)) if q else [],
        "processos": Processo.objects.filter(Q(numero__icontains=q) | Q(titulo__icontains=q) | Q(parte_autora__icontains=q) | Q(parte_re__icontains=q) | Q(descricao__icontains=q) | Q(tribunal__icontains=q)) if q else [],
        "materiais": Material.objects.filter(Q(nome__icontains=q) | Q(categoria__icontains=q) | Q(localizacao__icontains=q)) if q else [],
        "andamentos": Andamento.objects.filter(Q(descricao__icontains=q) | Q(responsavel__icontains=q)) if q else [],
    })

@login_required
def estoque(request):
    return render(request, "juridico/estoque.html", {"itens": Material.objects.all(), "movimentacoes": MovimentacaoEstoque.objects.select_related("material")[:10]})

def salvar(request, form_class, template, sucesso, instance=None):
    form = form_class(request.POST or None, instance=instance)
    if request.method == "POST" and form.is_valid(): form.save(); return redirect(sucesso)
    return render(request, "juridico/form.html", {"form": form, "titulo": template})

@login_required
def advogado_novo(request): return salvar(request, AdvogadoForm, "Novo advogado", "advogados")
@login_required
def advogado_editar(request, pk): return salvar(request, AdvogadoForm, "Editar advogado", "advogados", get_object_or_404(Advogado, pk=pk))
@login_required
def processo_novo(request): return salvar(request, ProcessoForm, "Novo processo", "processos")
@login_required
def processo_editar(request, pk): return salvar(request, ProcessoForm, "Editar processo", "processos", get_object_or_404(Processo, pk=pk))
@login_required
def andamento_novo(request):
    form = AndamentoForm(request.POST or None)

    if request.method == "POST" and form.is_valid():

        andamento = form.save()

        processo = andamento.processo

        novo_status = form.cleaned_data.get("novo_status")

        if novo_status:
            processo.status = novo_status
            processo.save(update_fields=["status", "atualizado_em"])

        messages.success(
            request,
            "Andamento registrado e status do processo atualizado."
        )

        return redirect(
            "processo_detalhe",
            pk=processo.pk
        )

    return render(
        request,
        "juridico/form.html",
        {
            "form": form,
            "titulo": "Novo andamento"
        }
    )
@login_required
def historico_novo(request): return salvar(request, HistoricoForm, "Novo item no histórico", "advogados")
@login_required
def material_novo(request): return salvar(request, MaterialForm, "Novo material", "estoque")
@login_required
def movimentacao_nova(request): return salvar(request, MovimentacaoForm, "Movimentar estoque", "estoque")

@login_required
def advogado_excluir(request, pk):
    item = get_object_or_404(Advogado, pk=pk)
    if request.method == "POST":
        try:
            nome = item.nome; item.delete(); messages.success(request, f"Advogado {nome} excluído com sucesso.")
        except ProtectedError:
            messages.error(request, "Este advogado possui processos vinculados. Transfira ou exclua os processos antes de excluí-lo.")
        return redirect("advogados")
    return render(request, "juridico/confirmar_exclusao.html", {"item": item, "tipo": "advogado", "cancelar": "advogados"})

@login_required
def processo_excluir(request, pk):
    item = get_object_or_404(Processo, pk=pk)
    if request.method == "POST":
        numero = item.numero
        for documento in item.documentos.all(): documento.arquivo.delete(save=False)
        item.delete(); messages.success(request, f"Processo {numero} excluído com sucesso."); return redirect("processos")
    return render(request, "juridico/confirmar_exclusao.html", {"item": item, "tipo": "processo", "cancelar": "processos"})

@login_required
def processo_importar(request):
    form = ImportarProcessoForm(request.POST or None, request.FILES or None)
    if request.method == "POST" and form.is_valid():
        arquivo = form.cleaned_data["arquivo"]
        try:
            texto = extrair_texto(arquivo)
            dados = analisar_processo(texto, form.cleaned_data.get("advogado"))
            processo = Processo.objects.create(**dados)
            DocumentoProcesso.objects.create(processo=processo, arquivo=arquivo, nome_original=arquivo.name, texto_extraido=texto)
            messages.success(request, f"Processo {processo.numero} importado e preenchido automaticamente.")
            return redirect("processo_detalhe", pk=processo.pk)
        except IntegrityError:
            form.add_error("arquivo", "Já existe um processo com o número identificado neste documento.")
        except (ValueError, OSError) as erro:
            form.add_error("arquivo", str(erro))
    return render(request, "juridico/importar_processo.html", {"form": form})

@login_required
def exportar(request, modulo, formato):
    if modulo not in EXPORTACOES or formato not in {"pdf", "word"}: return HttpResponse(status=404)
    titulo, cabecalhos, linha, consulta = EXPORTACOES[modulo]
    dados = [[str(v) for v in linha(obj)] for obj in consulta()]
    agora = timezone.localtime().strftime("%d/%m/%Y às %H:%M")
    if formato == "word":
        doc = Document(); doc.add_heading(f"Relatório de {titulo}", 0); doc.add_paragraph(f"Gerado em {agora}")
        table = doc.add_table(rows=1, cols=len(cabecalhos)); table.style = "Table Grid"
        for i, texto in enumerate(cabecalhos): table.rows[0].cells[i].text = texto
        for valores in dados:
            cells = table.add_row().cells
            for i, texto in enumerate(valores): cells[i].text = texto
        out = BytesIO(); doc.save(out); out.seek(0)
        response = HttpResponse(out.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f'attachment; filename="{modulo}.docx"'; return response
    out = BytesIO(); doc = SimpleDocTemplate(out, pagesize=landscape(A4), rightMargin=1*cm, leftMargin=1*cm, topMargin=1*cm, bottomMargin=1*cm)
    styles = getSampleStyleSheet(); elementos = [Paragraph(f"Relatório de {titulo}", styles["Title"]), Paragraph(f"Gerado em {agora}", styles["Normal"]), Spacer(1, .4*cm)]
    tabela = Table([cabecalhos] + dados, repeatRows=1)
    tabela.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#17375e")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("GRID", (0,0), (-1,-1), .5, colors.grey), ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"), ("FONTSIZE", (0,0), (-1,-1), 8), ("VALIGN", (0,0), (-1,-1), "TOP"), ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f3f6fa")])]))
    elementos.append(tabela); doc.build(elementos)
    response = HttpResponse(out.getvalue(), content_type="application/pdf"); response["Content-Disposition"] = f'attachment; filename="{modulo}.pdf"'; return response

@login_required
def peticoes_ia(request):
    itens = PeticaoIA.objects.select_related("processo", "advogado", "criado_por")
    q = request.GET.get("q", "").strip()
    if q:
        itens = itens.filter(Q(processo__numero__icontains=q) | Q(processo__titulo__icontains=q) | Q(advogado__nome__icontains=q) | Q(texto_gerado__icontains=q))
    return render(request, "juridico/peticoes_ia.html", {"itens": itens, "q": q})

@login_required
def peticao_ia_nova(request, processo_id):
    processo = get_object_or_404(Processo.objects.select_related("advogado"), pk=processo_id)
    form = PeticaoIAForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        peticao = form.save(commit=False)
        peticao.processo, peticao.advogado = processo, processo.advogado
        peticao.criado_por, peticao.modelo_utilizado = request.user, settings.OPENAI_MODEL
        try:
            peticao.texto_gerado = gerar_minuta_peticao(peticao, form.cleaned_data["usar_documentos"])
            peticao.save()
            messages.success(request, "Minuta gerada. Revise integralmente antes de qualquer utilização.")
            return redirect("peticao_ia_detalhe", pk=peticao.pk)
        except Exception as erro:
            messages.error(request, f"Não foi possível gerar a minuta: {erro}")
    return render(request, "juridico/peticao_ia_form.html", {"form": form, "processo": processo})

@login_required
def peticao_ia_detalhe(request, pk):
    peticao = get_object_or_404(PeticaoIA.objects.select_related("processo", "advogado", "criado_por"), pk=pk)
    return render(request, "juridico/peticao_ia_detalhe.html", {"peticao": peticao})

@login_required
def peticao_ia_revisar(request, pk):
    peticao = get_object_or_404(PeticaoIA, pk=pk)
    form = RevisarPeticaoForm(request.POST or None, instance=peticao)
    if request.method == "POST" and form.is_valid():
        peticao = form.save(commit=False)
        if peticao.status == "APROVADA": peticao.revisado_em = timezone.now()
        peticao.save(); messages.success(request, "Revisão salva com sucesso.")
        return redirect("peticao_ia_detalhe", pk=peticao.pk)
    return render(request, "juridico/peticao_ia_revisar.html", {"form": form, "peticao": peticao})

@login_required
def peticao_ia_excluir(request, pk):
    peticao = get_object_or_404(PeticaoIA, pk=pk)
    if request.method == "POST":
        peticao.delete(); messages.success(request, "Petição excluída com sucesso."); return redirect("peticoes_ia")
    return render(request, "juridico/confirmar_exclusao.html", {"item": peticao, "tipo": "petição", "cancelar": "peticoes_ia"})

def _nome_peticao(peticao):
    return "peticao_" + "".join(c if c.isalnum() or c in "-_" else "-" for c in peticao.processo.numero)

@login_required
def peticao_ia_word(request, pk):
    peticao = get_object_or_404(PeticaoIA.objects.select_related("processo"), pk=pk)
    documento = Document(); documento.add_heading(peticao.get_tipo_display(), 0)
    documento.add_paragraph(f"Processo: {peticao.processo.numero}")
    documento.add_paragraph("Minuta elaborada com auxílio de IA e sujeita à revisão jurídica.")
    for bloco in peticao.texto_gerado.splitlines():
        if bloco.strip(): documento.add_paragraph(bloco.strip())
    out = BytesIO(); documento.save(out); out.seek(0)
    response = HttpResponse(out.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="{_nome_peticao(peticao)}.docx"'; return response

@login_required
def peticao_ia_pdf(request, pk):
    peticao = get_object_or_404(PeticaoIA.objects.select_related("processo"), pk=pk)
    out = BytesIO(); documento = SimpleDocTemplate(out, pagesize=A4, rightMargin=2.5*cm, leftMargin=2.5*cm, topMargin=2.5*cm, bottomMargin=2.5*cm)
    estilos = getSampleStyleSheet(); elementos = [Paragraph(escape(peticao.get_tipo_display()), estilos["Title"]), Paragraph(f"Processo: {escape(peticao.processo.numero)}", estilos["Normal"]), Spacer(1, .7*cm)]
    for bloco in peticao.texto_gerado.splitlines():
        if bloco.strip(): elementos.extend([Paragraph(escape(bloco.strip()), estilos["BodyText"]), Spacer(1, .15*cm)])
        else: elementos.append(Spacer(1, .25*cm))
    documento.build(elementos); out.seek(0)
    response = HttpResponse(out.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{_nome_peticao(peticao)}.pdf"'; return response

# Import local intencionalmente no fim para manter a área de imports legível.
from django.db import models
