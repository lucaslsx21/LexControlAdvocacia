import re
import unicodedata
from datetime import datetime
from decimal import Decimal, InvalidOperation
from io import BytesIO

from django.utils import timezone
from docx import Document
from pypdf import PdfReader

from .models import Advogado


# ============================================================
# EXTRAÇÃO DO TEXTO DO DOCUMENTO
# ============================================================

def extrair_texto(arquivo):
    nome = arquivo.name.lower()

    conteudo = arquivo.read()
    arquivo.seek(0)

    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------
    if nome.endswith(".pdf"):
        try:
            leitor = PdfReader(BytesIO(conteudo))

            paginas = []

            for pagina in leitor.pages:
                texto_pagina = pagina.extract_text() or ""

                if texto_pagina.strip():
                    paginas.append(texto_pagina)

            texto = "\n".join(paginas)

        except Exception as erro:
            raise ValueError(
                f"Não foi possível ler o arquivo PDF: {erro}"
            )

    # --------------------------------------------------------
    # WORD
    # --------------------------------------------------------
    elif nome.endswith(".docx"):
        try:
            doc = Document(BytesIO(conteudo))

            partes = []

            # Parágrafos
            for paragrafo in doc.paragraphs:
                if paragrafo.text.strip():
                    partes.append(paragrafo.text)

            # Tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        if celula.text.strip():
                            partes.append(celula.text)

            texto = "\n".join(partes)

        except Exception as erro:
            raise ValueError(
                f"Não foi possível ler o arquivo DOCX: {erro}"
            )

    # --------------------------------------------------------
    # TXT
    # --------------------------------------------------------
    elif nome.endswith(".txt"):
        texto = conteudo.decode(
            "utf-8",
            errors="replace"
        )

    else:
        raise ValueError(
            "Formato de arquivo não suportado. "
            "Utilize PDF, DOCX ou TXT."
        )

    # --------------------------------------------------------
    # LIMPEZA DO TEXTO
    # --------------------------------------------------------

    texto = texto.replace("\xa0", " ")

    texto = re.sub(
        r"[ \t]+",
        " ",
        texto
    )

    texto = re.sub(
        r"\n{3,}",
        "\n\n",
        texto
    )

    texto = texto.strip()

    # Documento escaneado normalmente não terá texto suficiente
    if len(texto) < 30:
        raise ValueError(
            "Não foi encontrado texto suficiente no documento. "
            "O arquivo pode ser uma imagem digitalizada e precisar de OCR."
        )

    return texto


# ============================================================
# NORMALIZAÇÃO
# ============================================================

def _normalizar(texto):
    if not texto:
        return ""

    return "".join(
        caractere
        for caractere in unicodedata.normalize(
            "NFD",
            texto.upper()
        )
        if unicodedata.category(caractere) != "Mn"
    )


# ============================================================
# LOCALIZAR CAMPOS NO TEXTO
# ============================================================

def _campo(texto, rotulos, limite=180):
    """
    Procura informações utilizando vários possíveis rótulos.

    Exemplos:
        Autor: João da Silva
        Parte autora - João da Silva
        Ré: Empresa ABC Ltda.
        Requerida: Empresa ABC Ltda.

    Também aceita situações em que o valor aparece
    na linha imediatamente seguinte ao rótulo.
    """

    for rotulo in rotulos:

        rotulo_regex = re.escape(rotulo)

        # ----------------------------------------------------
        # FORMATO:
        #
        # Ré: Empresa ABC
        # Autor - João da Silva
        # ----------------------------------------------------

        padrao_mesma_linha = rf"""
            (?im)
            ^\s*
            {rotulo_regex}
            \s*
            [:\-–—]?
            \s*
            ([^\n\r]{{2,{limite}}})
            \s*$
        """

        resultado = re.search(
            padrao_mesma_linha,
            texto,
            re.VERBOSE | re.IGNORECASE | re.MULTILINE
        )

        if resultado:
            valor = resultado.group(1).strip(
                " .;:-–—"
            )

            if valor:
                return valor

        # ----------------------------------------------------
        # FORMATO:
        #
        # Parte ré:
        # Empresa ABC Ltda.
        # ----------------------------------------------------

        padrao_linha_seguinte = rf"""
            (?im)
            ^\s*
            {rotulo_regex}
            \s*
            [:\-–—]?
            \s*
            $
            \s*
            ([^\n\r]{{2,{limite}}})
        """

        resultado = re.search(
            padrao_linha_seguinte,
            texto,
            re.VERBOSE | re.IGNORECASE | re.MULTILINE
        )

        if resultado:
            valor = resultado.group(1).strip(
                " .;:-–—"
            )

            if valor:
                return valor

    # --------------------------------------------------------
    # FALLBACK
    #
    # Alguns PDFs perdem parcialmente a estrutura das linhas.
    # --------------------------------------------------------

    padrao_rotulos = "|".join(
        re.escape(rotulo)
        for rotulo in rotulos
    )

    resultado = re.search(
        rf"(?:{padrao_rotulos})"
        rf"\s*[:\-–—]?\s*"
        rf"([^\n\r]{{2,{limite}}})",
        texto,
        re.IGNORECASE
    )

    if resultado:
        return resultado.group(1).strip(
            " .;:-–—"
        )

    return ""


# ============================================================
# CONVERSÃO DE DATA
# ============================================================

def _data(valor):
    if not valor:
        return None

    resultado = re.search(
        r"(\d{2})[/-](\d{2})[/-](\d{4})"
        r"(?:\s+(?:às\s*)?(\d{1,2}):(\d{2}))?",
        valor,
        re.IGNORECASE
    )

    if not resultado:
        return None

    try:
        data = datetime(
            int(resultado.group(3)),
            int(resultado.group(2)),
            int(resultado.group(1)),
            int(resultado.group(4) or 0),
            int(resultado.group(5) or 0)
        )

        # Se possuir horário
        if resultado.group(4):
            if timezone.is_naive(data):
                data = timezone.make_aware(data)

            return data

        # Sem horário
        return data.date()

    except ValueError:
        return None


# ============================================================
# NÚMERO DO PROCESSO
# ============================================================

def _numero_processo(texto):
    """
    Procura primeiro pelo padrão CNJ.
    """

    # Exemplo:
    # 0000000-00.2026.8.08.0000

    resultado = re.search(
        r"\b"
        r"\d{7}"
        r"-"
        r"\d{2}"
        r"\."
        r"\d{4}"
        r"\."
        r"\d"
        r"\."
        r"\d{2}"
        r"\."
        r"\d{4}"
        r"\b",
        texto
    )

    if resultado:
        return resultado.group(0)

    # Outros formatos
    resultado = re.search(
        r"(?:PROCESSO|AUTOS)"
        r"(?:\s+N[º°O.]*)?"
        r"\s*[:\-]?\s*"
        r"([\w.\-/]{6,40})",
        texto,
        re.IGNORECASE
    )

    if resultado:
        return resultado.group(1).strip()

    return ""


# ============================================================
# IDENTIFICAÇÃO DA ÁREA JURÍDICA
# ============================================================

def _identificar_area(texto_normalizado):

    areas = [

        (
            "TRABALHO",
            [
                "DIREITO DO TRABALHO",
                "TRABALHISTA",
                "RECLAMACAO TRABALHISTA",
                "JUSTICA DO TRABALHO"
            ]
        ),

        (
            "TRIBUTARIO",
            [
                "DIREITO TRIBUTARIO",
                "TRIBUTARIO",
                "EXECUCAO FISCAL",
                "TRIBUTO"
            ]
        ),

        (
            "FAMILIA",
            [
                "DIREITO DE FAMILIA",
                "DIVORCIO",
                "ALIMENTOS",
                "GUARDA",
                "UNIAO ESTAVEL"
            ]
        ),

        (
            "CONSUMIDOR",
            [
                "DIREITO DO CONSUMIDOR",
                "CONSUMIDOR",
                "RELACAO DE CONSUMO",
                "CODIGO DE DEFESA DO CONSUMIDOR"
            ]
        ),

        (
            "PENAL",
            [
                "DIREITO PENAL",
                "PENAL",
                "CRIMINAL",
                "ACAO PENAL"
            ]
        ),

        (
            "CIVIL",
            [
                "DIREITO CIVIL",
                "CIVEL",
                "PROCEDIMENTO COMUM",
                "RESPONSABILIDADE CIVIL"
            ]
        ),
    ]

    for codigo, termos in areas:
        for termo in termos:
            if termo in texto_normalizado:
                return codigo

    return "CIVIL"


# ============================================================
# IDENTIFICAÇÃO DO STATUS
# ============================================================

def _identificar_status(texto_normalizado):

    if any(
        termo in texto_normalizado
        for termo in [
            "TRANSITADO EM JULGADO",
            "PROCESSO CONCLUIDO",
            "BAIXA DEFINITIVA",
            "ARQUIVADO DEFINITIVAMENTE"
        ]
    ):
        return "CONCLUIDO"

    if "INDEFERIDO" in texto_normalizado:
        return "INDEFERIDO"

    if "DEFERIDO" in texto_normalizado:
        return "DEFERIDO"

    if "SUSPENSO" in texto_normalizado:
        return "SUSPENSO"

    return "EM_ANDAMENTO"


# ============================================================
# CONVERSÃO DO VALOR DA CAUSA
# ============================================================

def _valor_monetario(valor_texto):

    if not valor_texto:
        return Decimal("0")

    resultado = re.search(
        r"(?:R\$\s*)?"
        r"([\d.]+,\d{2})",
        valor_texto
    )

    if not resultado:
        return Decimal("0")

    try:
        valor = resultado.group(1)

        valor = valor.replace(
            ".",
            ""
        ).replace(
            ",",
            "."
        )

        return Decimal(valor)

    except InvalidOperation:
        return Decimal("0")


# ============================================================
# LOCALIZAR ADVOGADO
# ============================================================

def _identificar_advogado(
    texto_normalizado,
    advogado_escolhido=None
):

    # Se o usuário selecionou manualmente,
    # essa opção tem prioridade.
    if advogado_escolhido:
        return advogado_escolhido

    for candidato in Advogado.objects.filter(
        ativo=True
    ):

        nome = _normalizar(
            candidato.nome
        )

        oab = _normalizar(
            candidato.oab
        )

        if nome and nome in texto_normalizado:
            return candidato

        if oab and oab in texto_normalizado:
            return candidato

    return None


# ============================================================
# ANÁLISE PRINCIPAL DO PROCESSO
# ============================================================

def analisar_processo(
    texto,
    advogado_escolhido=None
):

    normal = _normalizar(texto)

    # --------------------------------------------------------
    # NÚMERO DO PROCESSO
    # --------------------------------------------------------

    numero = _numero_processo(texto)

    # --------------------------------------------------------
    # ÁREA JURÍDICA
    # --------------------------------------------------------

    area = _identificar_area(normal)

    # --------------------------------------------------------
    # STATUS
    # --------------------------------------------------------

    status = _identificar_status(normal)

    # --------------------------------------------------------
    # PARTE AUTORA
    # --------------------------------------------------------

    autor = _campo(
        texto,
        [
            "parte autora",
            "parte autor",
            "autor",
            "autora",
            "requerente",
            "reclamante",
            "exequente",
            "demandante",
            "promovente",
            "polo ativo"
        ]
    )

    # --------------------------------------------------------
    # PARTE RÉ
    #
    # AQUI ESTAVA O PRINCIPAL PROBLEMA.
    # Agora reconhece "Ré:" também.
    # --------------------------------------------------------

    reu = _campo(
        texto,
        [
            "parte ré",
            "parte re",
            "réu",
            "reu",
            "ré",
            "re",
            "requerido",
            "requerida",
            "reclamado",
            "reclamada",
            "executado",
            "executada",
            "demandado",
            "demandada",
            "promovido",
            "promovida",
            "polo passivo"
        ]
    )

    # --------------------------------------------------------
    # TRIBUNAL
    # --------------------------------------------------------

    tribunal = _campo(
        texto,
        [
            "tribunal",
            "órgão julgador",
            "orgao julgador",
            "vara",
            "juízo",
            "juizo"
        ]
    )

    # --------------------------------------------------------
    # COMARCA
    # --------------------------------------------------------

    comarca = _campo(
        texto,
        [
            "comarca",
            "foro",
            "foro central"
        ]
    )

    # --------------------------------------------------------
    # CLASSE / TÍTULO
    # --------------------------------------------------------

    titulo = _campo(
        texto,
        [
            "classe judicial",
            "classe processual",
            "classe",
            "assunto",
            "tipo de ação",
            "tipo de acao"
        ]
    )

    if not titulo:
        titulo = (
            f"Processo {numero}"
            if numero
            else "Processo importado"
        )

    # --------------------------------------------------------
    # DATA DE DISTRIBUIÇÃO
    # --------------------------------------------------------

    distribuicao_texto = _campo(
        texto,
        [
            "data de distribuição",
            "data da distribuição",
            "data de distribuicao",
            "data da distribuicao",
            "distribuição",
            "distribuicao"
        ]
    )

    distribuicao = (
        _data(distribuicao_texto)
        or timezone.localdate()
    )

    # --------------------------------------------------------
    # PRAZO
    # --------------------------------------------------------

    prazo_texto = _campo(
        texto,
        [
            "prazo atual",
            "prazo",
            "data limite",
            "data-limite",
            "vencimento"
        ]
    )

    prazo = _data(prazo_texto)

    # --------------------------------------------------------
    # VALOR DA CAUSA
    # --------------------------------------------------------

    valor_texto = _campo(
        texto,
        [
            "valor da causa",
            "valor atribuído à causa",
            "valor atribuido a causa"
        ]
    )

    valor = _valor_monetario(
        valor_texto
    )

    # --------------------------------------------------------
    # ADVOGADO
    # --------------------------------------------------------

    advogado = _identificar_advogado(
        normal,
        advogado_escolhido
    )

    # --------------------------------------------------------
    # CAMPOS OBRIGATÓRIOS
    # --------------------------------------------------------

    obrigatorios = []

    if not numero:
        obrigatorios.append(
            "número do processo"
        )

    if not autor:
        obrigatorios.append(
            "parte autora"
        )

    if not reu:
        obrigatorios.append(
            "parte ré"
        )

    if not advogado:
        obrigatorios.append(
            "advogado responsável"
        )

    if obrigatorios:

        campos = ", ".join(
            obrigatorios
        )

        raise ValueError(
            "Não foi possível identificar: "
            + campos
            + ". Informe esses dados no documento "
              "ou selecione o advogado no envio."
        )

    # --------------------------------------------------------
    # PREPARAÇÃO DO PRAZO
    # --------------------------------------------------------

    # Seu model aparentemente utiliza DateTimeField.
    # Por isso mantemos somente datetime aqui.
    prazo_atual = (
        prazo
        if isinstance(
            prazo,
            datetime
        )
        else None
    )

    # --------------------------------------------------------
    # RETORNO PARA CRIAÇÃO DO PROCESSO
    # --------------------------------------------------------

    return {
        "numero": numero[:40],

        "titulo": titulo[:180],

        "area": area,

        "status": status,

        "advogado": advogado,

        "parte_autora": autor[:180],

        "parte_re": reu[:180],

        "tribunal": tribunal[:150],

        "comarca": comarca[:120],

        "data_distribuicao": distribuicao,

        "prazo_atual": prazo_atual,

        "valor_causa": valor,

        "descricao": texto[:10000],
    }